# -*- coding: utf-8 -*-

from docx import Document

#This script takes information about the case and puts into a properly formatted word document so that the person writing it does not have to do the tedious work of writing in almost 90 pieces of information by hand. A robot should be doing that, not a human!

#upload the blank template that conforms to style guide with tagged text to be replaced
document = Document('Brief Template.docx')

#information specific to the case, and who is doing it. 
submittedORargued = 'argued'
attorney = 'David Billingsley'
title = 'Staff Attorney'
client = 	'CLIENT NAME'
email = "email"
supervisor = 'SUPERVISOR NAME'
judgmentDate = 'DATE'
pleaORtrial = 'bench trial'

if ('trial' in pleaORtrial):
  tp = "T."
else: tp = "P."

#no punctuation at the end for these.
countStatement = 'two counts of Robbery in the Second Degree, one each under Penal Law § 160.10(1) and Penal Law § 160.10(2)(a), one count of Strangulation in the Second Degree, 121.12, and four counts of Grand Larceny in the Fourth Degree, Penal Law § 155.30(4)'
clientShortName = 'Mr. LASTNAME'
sentence = '25 years to life in prison on each of the robbery and strangulation charges, and 2 to 4 years in prison on the grand larceny charges, all to be served concurrently'
justiceStatement = 'Justice NAME presided over the suppression hearing. Justice NAME presided over the bench trial and sentencing'
status = "serving his term of imprisonment"
statusStatement = ""
sentenceDone = False

if sentenceDone:
  statusStatement =  clientShortName + "has been released following the completion of his sentence"
else: statusStatement = clientShortName + ' is currently ' + status

indictment = 'xxxx-xxxx'
county = 'New York'
appealDate = 'DATE'
recordDate = 'DATE'
serveDate = 'August ____, 2019'
sentenceDate = 'July ___, 2016'
pleaOrTrialDate = 'DATE - DATE'
term = "November 2019"
termDate = "September 3, 2019"
hearingDate = "DATE - DATE" 

fileName = client + " " + indictment + ' Draft'

mainbodycount = 0
tablecount = 0
footercount = 0

# stole this method from here https://stackoverflow.com/questions/34779724/python-docx-replace-string-in-paragraph-while-keeping-style, with a slight edit so that it recognizes the § character (that's what decode("utf8") is.)
def replace_string(oldText, newText):
    
    #replaces tagged text in the main body
    for p in document.paragraphs:
        if oldText in p.text:
            global mainbodycount
            print 'Main body replacement!'
            mainbodycount += 1
            text = p.text.replace(oldText, newText.decode("utf8"))
            style = p.style #not sure these style lines are necessary
            p.text = text
            p.style = style 
            
    
    #replaces tagged text in the caption, which is part of a table
    for t in document.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if oldText in p.text:
                        global tablecount
                        print 'Table Replacement!'
                        tablecount += 1
                        text = p.text.replace(oldText, newText.decode("utf8"))
                        style = p.style
                        p.text = text
                        p.style = style

    #need to add a replace loop for the header and footer...not sure if possible though
    
    for s in document.sections:
      for p in s.footer.paragraphs:
        if oldText in p.text:
          global footercount
          print 'Footer Replacement!'
          footercount += 1
          text = p.text.replace(oldText, newText.decode("utf8"))
          
          if pleaORtrial == 'trial':
            text = text.replace('\"P.\"', '\"T.\"')
            text = text.replace("the guilty plea", "trial proceedings")

          style = p.style
          p.text = text
          p.style = style


    return 1

#replace_string is a little weird in that it does not preserve formatting, so what you'll have to do here, is change the fields in the document to have a caps, e.g., [APPELLANT-CAPS], when it should be in caps in the document. And then do a replace statement that capitalizes it, e.g., client.upper()

replace_string('[APPELLANT]', client)
replace_string('[APPELLANT-CAPS]', client.upper())
replace_string('[SUPERVISOR]', supervisor)
replace_string('[SUPERVISOR-CAPS]', supervisor.upper())
replace_string('[JUDGMENTDATE]', judgmentDate)
replace_string('[GUILTY PLEA/TRIAL]', pleaORtrial)
replace_string('[MR/MS LAST NAME]', clientShortName)
replace_string('[SENTENCE]', sentence)
replace_string('[ATTORNEY]', attorney)
replace_string('[ATTORNEY-CAPS]', attorney.upper())
replace_string('[TITLE]', title)
replace_string('[EMAIL]', email)
replace_string('[INDICTMENT]', indictment)
replace_string('[COUNTY]', county)
replace_string('[COUNTS]', countStatement)
replace_string('[JUSTICE STATEMENT]', justiceStatement)
replace_string('[STATUS]', statusStatement)
replace_string('[TERM]', term)
replace_string('[APPEAL DATE]', appealDate)
replace_string('[RECORD DATE]', recordDate)
replace_string('[SERVE DATE]', serveDate)
replace_string('[TERM DATE]', termDate)
replace_string('[PLEA-TRIAL DATE]', pleaOrTrialDate)
replace_string("[HEARING-DATE]", hearingDate)
replace_string("[T-P]", tp)
replace_string("[SENTENCE DATE]", sentenceDate)

print "Total replacements: " + str(mainbodycount + tablecount + footercount)

document.save(fileName + '.docx')

#To do: Header and Footer Replacements
#To do: Asa will do the web interface, I will modify accordingly.
#To do: Expand to other types of filings like 440.
#To do: How to handle multiple cases in one brief?